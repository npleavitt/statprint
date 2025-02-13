from docx import Document
from docx.shared import Inches, Pt
from fpdf import FPDF
import pandas as pd

class StatPrint:
    def __init__(self, filename="report", doc_type="word", title="Report", table_theme=None):
        # Initialize with default filename, document type, and title
        self.filename = filename
        self.doc_type = doc_type
        self.title = title
        self.content = []
        self.graph_count = 0  # To create unique filenames for graphs
        
        # Set a default table theme if one isn't provided.
        # Colors are given as hex strings (without '#').
        # 'table_width' is in twips (1/20 of a point, commonly used in Word's XML) or as a string;
        # if None, the table will use its default width.
        if table_theme is None:
            self.table_theme = {
                'header_bg_color': 'D9D9D9',
                'row_bg_color_even': 'F2F2F2',
                'row_bg_color_odd': None,  # Leave odd rows with no fill
                'table_width': None  # e.g., '5000' (dxa) to set a fixed width, or None to auto-fit
            }
        else:
            self.table_theme = table_theme

    def add_cover_page(self, title, subtitle=None, author=None, date=None):
        """
        Add a cover page to the report. This will be inserted at the beginning.
        """
        cover_content = {
            'title': title,
            'subtitle': subtitle,
            'author': author,
            'date': date
        }
        # Insert at the beginning of content
        self.content.insert(0, ('cover', cover_content))

    def add_heading(self, heading):
        """Add a heading to the report."""
        self.content.append(('heading', heading))

    def add_table(self, data, custom_headers=None, indent_rows=None):
        """
        Add a table to the report.
        
        Parameters:
          - data: a pandas DataFrame or Series.
          - custom_headers: Optional list of column names to override the default.
          - indent_rows: Optional list of row indices (starting at 0 for the first data row)
                         that should have their first cell indented.
        """
        if indent_rows is None:
            indent_rows = []

        if isinstance(data, pd.Series):
            # For Series (often from value_counts), reset the index so that you have two columns.
            df = data.reset_index()
            if custom_headers is None:
                if data.name is not None:
                    custom_headers = [data.name, "Count"]
                else:
                    custom_headers = ["Value", "Count"]
            df.columns = custom_headers
        else:
            df = data.copy()
            if custom_headers is not None:
                df.columns = custom_headers

        # Store the table with additional styling options.
        self.content.append(('table', {'df': df, 'custom_headers': custom_headers, 'indent_rows': indent_rows}))

    def add_graph(self, graph, filename=None):
        """
        Save the graph (a matplotlib figure) to a unique file and add its filename to the report.
        
        Parameters:
          - graph: a matplotlib figure object.
          - filename: Optional filename base; if provided the counter is prepended.
        """
        if filename is None:
            filename = f"graph_{self.graph_count}.png"
        else:
            filename = f"{self.graph_count}_{filename}"
        self.graph_count += 1
        graph.savefig(filename, format='png')
        self.content.append(('graph', filename))

    def generate_report(self):
        if self.doc_type == 'word':
            self.generate_word_report()
        elif self.doc_type == 'pdf':
            self.generate_pdf_report()

    def _set_cell_background(self, cell, color):
        """
        Helper method to set the background shading for a cell.
        :param cell: A docx TableCell object.
        :param color: A hex color string (without the '#' symbol).
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tcPr = cell._tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            cell._tc.insert(0, tcPr)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

    def _apply_word_table_style(self, table, indent_rows):
        """
        Helper method to modify a docx table:
          - Applies header styling (bold text and background color).
          - Sets alternating row shading for data rows.
          - Applies left indentation for specified rows.
          - Optionally sets a fixed table width.
        """
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls, qn

        # Process header row: bold text and header background.
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            # Set header background color.
            if self.table_theme.get('header_bg_color'):
                self._set_cell_background(cell, self.table_theme['header_bg_color'])

        # Process data rows: alternating shading and indent if required.
        for i, row in enumerate(table.rows[1:]):  # skip header row
            # Determine fill color based on row index.
            if i % 2 == 0 and self.table_theme.get('row_bg_color_even'):
                fill_color = self.table_theme['row_bg_color_even']
            elif i % 2 == 1 and self.table_theme.get('row_bg_color_odd'):
                fill_color = self.table_theme['row_bg_color_odd']
            else:
                fill_color = None

            for cell in row.cells:
                if fill_color:
                    self._set_cell_background(cell, fill_color)

            # Apply left indentation if this row is marked for indenting.
            if i in indent_rows:
                first_cell = row.cells[0]
                for paragraph in first_cell.paragraphs:
                    paragraph.paragraph_format.left_indent = Inches(0.25)

        # Optionally, set a fixed table width.
        if self.table_theme.get('table_width'):
            tbl = table._element
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:w'), str(self.table_theme['table_width']))
            tblW.set(qn('w:type'), 'dxa')

        # Adjust borders to remove vertical lines (keeping horizontal lines).
        tblBorders = parse_xml(
            r'<w:tblBorders %s>'
            r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:left w:val="nil"/>'
            r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:right w:val="nil"/>'
            r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:insideV w:val="nil"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr.append(tblBorders)

    def generate_word_report(self):
        doc = Document()

        # If a cover page was added, process it first.
        # Check if the first content element is a cover.
        if self.content and self.content[0][0] == 'cover':
            cover = self.content.pop(0)[1]
            # Create a cover page with the provided details.
            doc.add_paragraph()  # add some spacing at the top
            title_par = doc.add_heading(cover.get('title', self.title), level=0)
            if cover.get('subtitle'):
                sub_par = doc.add_paragraph(cover.get('subtitle'))
            if cover.get('author'):
                doc.add_paragraph("Author: " + cover.get('author'))
            if cover.get('date'):
                doc.add_paragraph("Date: " + cover.get('date'))
            # Add a page break after the cover page.
            doc.add_page_break()

        # Add the report title.
        doc.add_heading(self.title, 0)

        for content_type, content in self.content:
            if content_type == 'heading':
                doc.add_heading(content, level=1)
            elif content_type == 'table':
                table_info = content
                df = table_info['df']
                indent_rows = table_info.get('indent_rows', [])
                # Create a table with one header row and as many columns as the DataFrame has.
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Table Grid'

                # Add header cells.
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(df.columns):
                    hdr_cells[i].text = str(col_name)
                    # Bold header text and apply header background via _apply_word_table_style.
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Add data rows.
                for idx, row in enumerate(df.itertuples(index=False, name=None)):
                    row_cells = table.add_row().cells
                    for j, value in enumerate(row):
                        row_cells[j].text = str(value)
                        # Apply indentation if needed.
                        if j == 0 and idx in indent_rows:
                            for paragraph in row_cells[j].paragraphs:
                                paragraph.paragraph_format.left_indent = Inches(0.25)

                # Apply publication-ready styling (header styling, alternating row shading, etc.).
                self._apply_word_table_style(table, indent_rows)
            elif content_type == 'graph':
                doc.add_paragraph()  # add spacing before image
                doc.add_picture(content, width=Inches(6))  # adjust width as needed

        doc.save(f"{self.filename}.docx")
        print(f"Report saved as {self.filename}.docx")

    def generate_pdf_report(self):
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, txt=self.title, ln=True, align='C')
        pdf.ln(5)

        pdf.set_font("Arial", '', 12)
        for content_type, content in self.content:
            if content_type == 'heading':
                pdf.ln(10)
                pdf.set_font("Arial", 'B', 14)
                pdf.cell(0, 10, txt=content, ln=True)
                pdf.set_font("Arial", '', 12)
            elif content_type == 'table':
                table_info = content
                df = table_info['df']
                indent_rows = table_info.get('indent_rows', [])

                # Calculate effective page width.
                effective_width = pdf.w - 2 * pdf.l_margin
                col_width = effective_width / len(df.columns)
                
                # Header row (bold, no cell borders; draw a horizontal line below).
                pdf.set_font("Arial", 'B', 10)
                for col in df.columns:
                    pdf.cell(col_width, 10, str(col), border=0, align='C')
                pdf.ln(10)
                # Draw horizontal line below header.
                pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + effective_width, pdf.get_y())

                # Data rows.
                pdf.set_font("Arial", '', 10)
                for i, row in enumerate(df.itertuples(index=False, name=None)):
                    for j, value in enumerate(row):
                        text = str(value)
                        if j == 0 and i in indent_rows:
                            text = "    " + text  # Indent if needed.
                        pdf.cell(col_width, 10, text, border=0)
                    pdf.ln(10)
                    # Draw horizontal line after each row.
                    pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + effective_width, pdf.get_y())
            elif content_type == 'graph':
                pdf.ln(10)
                pdf.image(content, x=None, y=None, w=150)  # Adjust width as necessary

        pdf.output(f"{self.filename}.pdf")
        print(f"Report saved as {self.filename}.pdf")
